from docx import Document
from config import path_config
import content
import util
import re


def check_font_size(run):
    '''检查字体大小
    '''
    try:
        return run.font.size.pt
    except AttributeError:
        return 'None'


def get_text(paragraphs):
    '''检查处理正文信息
    '''
    to_write = ""
    for paragraph in paragraphs:
        par_str = paragraph.text
        par_runs = paragraph.runs
        if len(par_runs):
            if len(par_str) and check_font_size(par_runs[0]) == 'None':
                par_str = content.tidy_para_text(par_str)
                if par_str[0].islower():
                    to_write = to_write[:-1] + ' ' + par_str + '\n'
                else:
                    to_write = to_write + par_str + '\n'
    return to_write


if __name__ == "__main__":
    pc = path_config()
    wap = pc.get_wap()
    tap = pc.get_tap()
    fp, dp = util.retrieve_input_path(wap, 'docx')
    for file_name in fp:
        o_f = util.retrieve_output_path(tap, file_name)
        if not o_f:
            print('Conversion cancelled')
            continue
        doc = Document(util.full_path(wap, file_name))
        util.save_txt(util.full_path(tap, o_f), get_text(doc.paragraphs))
    
    # util.save_txt("F:\\PerStudy\\Desktop\\ParticipleDocs\\parseEEDoc\\articles-txt\\Debtholders' Demand for Conservatism Evidence from Changes in Directors' Fiduciary Duties.txt", get_text(doc.paragraphs))
