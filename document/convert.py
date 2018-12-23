import os
import logging
import winerror
import comtypes.client
from win32com.client.dynamic import Dispatch, ERRORS_BAD_CONTEXT

from .utils.util import ModefyPath, save_txt
from .utils.content import Content, strQ2B
from .config.Config import WordFormat
logging.getLogger().setLevel(logging.INFO)
ERRORS_BAD_CONTEXT.append(winerror.E_NOTIMPL)
mp = ModefyPath()
ct = Content()
wf = WordFormat()
wdFormatPDF = wf.get_format('PDF')

in_file_type = {
    'pdf_to_word': 'pdf',
    'word_to_pdf': 'docx',
    'word_to_txt': 'docx',
    'get_word': 'docx'
}
out_file_type = {
    'pdf_to_word': 'docx',
    'word_to_pdf': 'pdf',
    'word_to_txt': 'txt'
}

def word_to_pdf(word_file_path, pdf_file_path):
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(word_file_path)
    doc.SaveAs(pdf_file_path, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()


def pdf_to_word(pdf_file_path, word_file_path):
    avDoc = Dispatch("AcroExch.AVDoc")
    ret = avDoc.Open(pdf_file_path, pdf_file_path)
    assert(ret)
    pdDoc = avDoc.GetPDDoc()
    jsObject = pdDoc.GetJSObject()
    jsObject.SaveAs(word_file_path, "com.adobe.acrobat.docx")
    pdDoc.Close()
    avDoc.Close(True)


'''
word to txt need docx content
'''
from docx import Document


def get_docx_fontsize(run):
    '''检查字体大小
    '''
    try:
        return run.font.size.pt
    except AttributeError:
        return 'None'

def is_paragraph(article_type, *args):
    '''
    @args[0]: run
    @args[1]: paragraph
    '''
    if article_type == 'ICCAE':
        if get_docx_fontsize(args[0]) == 'None':
            return True
        else:
            return False
    # if article_type == 'PSR':
    #     pdfreader()

def get_docx_content(article_type, paragraphs):
    '''检查处理正文信息
    '''
    content = ""
    for paragraph in paragraphs:
        par_str = strQ2B(paragraph.text)
        par_runs = paragraph.runs
        # par_runs == 0 => par_str == 0
        if len(par_runs) and is_paragraph(article_type, par_runs):
            par_str = ct.tidy_paragraph(article_type, par_str)
            if not par_str:
                continue
            if par_str[0].islower():
                content = content[:-1] + ' ' + par_str + '\n'
            else:
                content = content + par_str + '\n'
    return content


def word_to_txt(word_file_path, txt_file_path):
    doc = Document(word_file_path)
    save_txt(txt_file_path, get_docx_content('ICCAE', doc.paragraphs))


def get_word(word_file_path):
    doc = Document(word_file_path)
    return get_docx_content('ICCAE', doc.paragraphs)


'''
pdf to txt need pdf content
'''
# def pdf_to_txt(pdf_file_path, txt_file_path):


'''
converter
'''


# def file_converter(input_file, output_file, function='pdf_to_word'):
def folder_converter(input_folder_relative, function='get_word', output_folder_relative=''):
    file_list = mp.load_file(input_folder_relative, in_file_type[function])
    if function.startswith("get"):
        document = {}
        for file_name in file_list:
            input_file = mp.get_full_path(input_folder_relative, file_name)
            document[file_name] = eval(function)(input_file)
        return document
    else:
        for file_name in file_list:
            o_f = mp.retrieve_file(output_folder_relative,
                                file_name, out_file_type[function])
            if not o_f:
                logging.info('Conversion cancelled')
                continue
            input_file = mp.get_full_path(input_folder_relative, file_name)
            output_file = mp.get_full_path(output_folder_relative, o_f)
            eval(function)(input_file, output_file)


if __name__ == '__main__':
    print(folder_converter('./articles/word', 'get_word'))


# from config import path_config
# import content
# import util
# import re


# def check_font_size(run):
#     '''检查字体大小
#     '''
#     try:
#         return run.font.size.pt
#     except AttributeError:
#         return 'None'


# def get_text(paragraphs):
#     '''检查处理正文信息
#     '''
#     to_write = ""
#     for paragraph in paragraphs:
#         par_str = paragraph.text
#         par_runs = paragraph.runs
#         if len(par_runs):
#             if len(par_str) and check_font_size(par_runs[0]) == 'None':
#                 par_str = content.tidy_para_text(par_str)
#                 if par_str[0].islower():
#                     to_write = to_write[:-1] + ' ' + par_str + '\n'
#                 else:
#                     to_write = to_write + par_str + '\n'
#     return to_write





# if __name__ == "__main__":
#     word_to_txt()
